import json
import logging
from pathlib import Path
from typing import Dict, Any

logger = logging.getLogger(__name__)

# Fallbacks for optional export libraries
HAS_FPDF = False
try:
    from fpdf import FPDF
    HAS_FPDF = True
except ImportError:
    logger.warning("fpdf2 package not found. PDF exports will be text/HTML representation only.")

HAS_DOCX = False
try:
    import docx
    HAS_DOCX = True
except ImportError:
    logger.warning("python-docx package not found. DOCX exports will fall back to text format.")

class DocumentExporter:
    @staticmethod
    def to_markdown(title: str, markdown_content: str) -> str:
        """Format notes or transcripts into standard markdown."""
        return f"# {title}\n\n{markdown_content}"

    @staticmethod
    def to_html(title: str, markdown_content: str) -> str:
        """Converts basic markdown content to styled HTML."""
        # Simple markdown to HTML parser (handles headings and bullet points)
        html_lines = []
        for line in markdown_content.split('\n'):
            line = line.strip()
            if not line:
                html_lines.append("<br/>")
            elif line.startswith('### '):
                html_lines.append(f"<h3>{line[4:]}</h3>")
            elif line.startswith('## '):
                html_lines.append(f"<h2>{line[3:]}</h2>")
            elif line.startswith('# '):
                html_lines.append(f"<h1>{line[2:]}</h1>")
            elif line.startswith('- ') or line.startswith('* '):
                html_lines.append(f"<li>{line[2:]}</li>")
            else:
                html_lines.append(f"<p>{line}</p>")

        html_body = "\n".join(html_lines)

        # Wrap in a gorgeous dark glassmorphic styling
        styled_html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>{title}</title>
    <style>
        body {{
            background: #0f0f13;
            color: #e2e2e9;
            font-family: 'Inter', -apple-system, sans-serif;
            padding: 40px;
            max-width: 800px;
            margin: 0 auto;
            line-height: 1.6;
        }}
        h1 {{ color: #a78bfa; border-bottom: 1px solid #2d2d3d; padding-bottom: 10px; }}
        h2 {{ color: #60a5fa; margin-top: 30px; }}
        h3 {{ color: #34d399; }}
        p {{ margin-bottom: 15px; color: #cbd5e1; }}
        li {{ margin-left: 20px; margin-bottom: 5px; color: #cbd5e1; }}
        br {{ content: ""; display: block; margin: 10px 0; }}
    </style>
</head>
<body>
    {html_body}
</body>
</html>
"""
        return styled_html

    @staticmethod
    def to_docx(title: str, markdown_content: str, output_path: Path) -> Path:
        """Exports markdown content to a Microsoft Word document (.docx)."""
        if not HAS_DOCX:
            # Fallback: save as raw text with docx suffix
            output_path.write_text(f"{title}\n\n{markdown_content}", encoding='utf-8')
            return output_path

        doc = docx.Document()
        doc.add_heading(title, 0)

        for line in markdown_content.split('\n'):
            line = line.strip()
            if not line:
                continue
            if line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            else:
                doc.add_paragraph(line)

        doc.save(str(output_path))
        return output_path

    @staticmethod
    def to_pdf(title: str, markdown_content: str, output_path: Path) -> Path:
        """Exports markdown content to a PDF document using FPDF2."""
        if not HAS_FPDF:
            # Fallback: write as HTML file with PDF extension or write text
            output_path.write_text(f"{title}\n\n{markdown_content}", encoding='utf-8')
            return output_path

        class PDF(FPDF):
            def header(self):
                self.set_font('Helvetica', 'B', 8)
                self.set_text_color(150, 150, 150)
                self.cell(0, 10, 'Local Media Brain Export', 0, 0, 'R')
                self.ln(15)
                
            def footer(self):
                self.set_y(-15)
                self.set_font('Helvetica', 'I', 8)
                self.set_text_color(150, 150, 150)
                self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')

        pdf = PDF()
        pdf.add_page()
        
        # Title
        pdf.set_font('Helvetica', 'B', 16)
        pdf.set_text_color(120, 100, 220)
        # Convert text to latin-1 to avoid fpdf character encoding limits
        pdf.cell(0, 10, title.encode('latin-1', 'ignore').decode('latin-1'), ln=True)
        pdf.ln(5)
        
        # Body
        pdf.set_font('Helvetica', '', 10)
        pdf.set_text_color(50, 50, 50)
        
        for line in markdown_content.split('\n'):
            line = line.strip()
            if not line:
                pdf.ln(4)
                continue
                
            clean_line = line.encode('latin-1', 'ignore').decode('latin-1')
            
            if line.startswith('### '):
                pdf.set_font('Helvetica', 'B', 11)
                pdf.set_text_color(20, 120, 80)
                pdf.multi_cell(0, 8, clean_line[4:])
                pdf.set_font('Helvetica', '', 10)
                pdf.set_text_color(50, 50, 50)
            elif line.startswith('## '):
                pdf.set_font('Helvetica', 'B', 13)
                pdf.set_text_color(30, 80, 150)
                pdf.multi_cell(0, 8, clean_line[3:])
                pdf.set_font('Helvetica', '', 10)
                pdf.set_text_color(50, 50, 50)
            elif line.startswith('# '):
                pdf.set_font('Helvetica', 'B', 15)
                pdf.set_text_color(80, 50, 150)
                pdf.multi_cell(0, 10, clean_line[2:])
                pdf.set_font('Helvetica', '', 10)
                pdf.set_text_color(50, 50, 50)
            elif line.startswith('- ') or line.startswith('* '):
                pdf.multi_cell(0, 6, f"  o  {clean_line[2:]}")
            else:
                pdf.multi_cell(0, 6, clean_line)
                
        pdf.output(str(output_path))
        return output_path

    @staticmethod
    def to_json(title: str, markdown_content: str) -> str:
        """Pack title and content as structured JSON."""
        return json.dumps({
            "title": title,
            "content": markdown_content
        }, indent=2)
